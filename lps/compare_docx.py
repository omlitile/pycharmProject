from langchain_community.document_loaders import DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import InMemoryVectorStore
from langchain_community.chains.retrieval_qa import RetrievalQAChain
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
import difflib
import docx
from deepdiff import DeepDiff
import pandas as pd
import env
import os

# 初始化OpenAI模型和Embeddings
llm = OpenAI(
        model="gpt-4-turbo",
        temperature=0.1,  # 严格遵循提示要求
        top_p=0.85,  # 保留高概率差异点
        max_tokens=500,  # 限制单次输出长度
        frequency_penalty=0.8,  # 防止重复报告相同差异
        presence_penalty = 0.4,  # 避免遗漏新条款
        api_key=os.getenv("OPENAI_API_KEY")
    )
embeddings = OpenAIEmbeddings(
    model="text-embedding-3-large",  # 指定大模型
    api_key=os.getenv("OPENAI_API_KEY")       # API密钥
)

def load_and_parse(file_path):
    """加载并解析docx文件，返回文本和表格数据"""
    doc = docx.Document(file_path)
    content = []
    tables = []

    # 提取段落文本（保留格式信息）
    for para in doc.paragraphs:
        text = {
            "text": para.text.strip(),
            "style": para.style.name,
            "bold": any(run.bold for run in para.runs),
            "italic": any(run.italic for run in para.runs)
        }
        if text["text"]: content.append(text)

    # 提取表格数据（保留行列结构）
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_data.append({
                "row": row.row_idx,
                "cells": [{"text": cell, "col_span": cell.col_span} for cell in cells]
            })
        if table_data: tables.append(table_data)

    return {"text": content, "tables": tables}


def semantic_chunker(text_blocks, chunk_size=512):
    """基于语义的分块策略（保留上下文关联）"""
    text = " ".join([f"[{blk['style']}] {blk['text']}" for blk in text_blocks])
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=100)
    return splitter.split_text(text)


def create_diff_chain(vector_store):
    """创建差异分析链，结合语义和结构对比"""
    prompt_template = """基于以下上下文差异分析，请以结构化JSON格式输出差异报告：
[原始内容]
{original}

[新内容]
{new}

[要求]
1. 识别新增/删除/修改的文本内容
2. 标注表格结构变化（行列增减、单元格内容变更）
3. 检测格式变化（加粗/斜体等样式）
4. 输出包含置信度的差异条目

返回格式：
[
    {
        "type": "text|table|format",
        "change_type": "新增|删除|修改",
        "original": "内容",
        "new": "内容",
        "position": "段落索引",
        "confidence": 0.85
    }
]"""

    return MultiRetrievalQAChain.from_llm(
        llm=llm,
        chain_type="stuff",
        retriever=vector_store.as_retriever(search_kwargs={"k": 3}),
        prompt_template=prompt_template,
        return_source_documents=True
    )


def compare_documents(file1, file2):
    """主对比函数，整合多维度差异检测"""
    # 加载解析网页
    data1 = load_and_parse(file1)
    data2 = load_and_parse(file2)

    # 语义向量化
    text1_chunks = semantic_chunker(data1["text"])
    text2_chunks = semantic_chunker(data2["text"])

    vector_store1 = InMemoryVectorStore.from_texts(
        text1_chunks,
        embedding=embeddings,
        prefix="file1: "
    )

    vector_store2 = InMemoryVectorStore.from_texts(
        text2_chunks,
        embedding=embeddings,
        prefix="file2: "
    )

    # 文本差异分析
    qa_chain = create_diff_chain(vector_store1)
    context_diff = f"文件1内容：{' '.join(text1_chunks)}\n文件2内容：{' '.join(text2_chunks)}"
    text_diff = qa_chain.run(context_diff=context_diff)

    # # 表格结构对比
    # table_diff = []
    # for t1, t2 in zip(data1["tables"], data2["tables"]):
    #     if len(t1) != len(t2):
    #         table_diff.append({"type": "table", "change": "行数变更",
    #                            "original": len(t1), "new": len(t2)})
    #     else:
    #         for r1, r2 in zip(t1, t2):
    #             if r1["row"] != r2["row"]:
    #                 table_diff.append({"type": "table", "change": "行序变化",
    #                                    "row": r1["row"], "new": r2["row"]})
    #             if r1["cells"] != r2["cells"]:
    #                 table_diff.append({"type": "table", "change": "单元格变更",
    #                                    "row": r1["row"],
    #                                    "details": deepdiff.DeepDiff(r1["cells"], r2["cells"])})

    # 格式差异检测
    format_diff = []
    for b1, b2 in zip(data1["text"], data2["text"]):
        if b1["style"] != b2.get("style", ""):
            format_diff.append({"type": "format", "change": "样式变更",
                                "text": b1["text"], "original": b1["style"],
                                "new": b2.get("style", "")})

    return {
        "semantic_diff": text_diff,
        # "table_diff": table_diff,
        "format_diff": format_diff
    }


# 使用示例
if __name__ == "__main__":
    file1 = "demo1.docx"
    file2 = "demo2.docx"

    report = compare_documents(file1, file2)

    # 生成可视化报告
    print("=== 差异分析报告 ===")
    print(f"语义差异条目: {len(report['semantic_diff'])}")
    # print(f"表格结构变化: {len(report['table_diff'])}")
    print(f"格式变更次数: {len(report['format_diff'])}")

    # 输出关键差异示例
    print("\n示例差异:")
    for diff in report["semantic_diff"][:3]:
        print(f"[{diff['type'].upper()}] {diff['change_type']} - 置信度{diff['confidence']:.2f}")
        print(f"原文: {diff['original']}\n新文: {diff['new']}\n")